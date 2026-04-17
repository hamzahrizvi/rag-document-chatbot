import os
import time
import shutil
import tempfile
import streamlit as st
from langchain_core.documents import Document
from langchain_ollama import OllamaLLM, OllamaEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from sentence_transformers import CrossEncoder
from pypdf import PdfReader
from docx import Document as DocxDocument

# --- CONFIG ---
VECTOR_STORE_PATH = "vector_store"
EMBED_MODEL = "nomic-embed-text"

st.set_page_config(page_title="RAG Chatbot", layout="wide")
st.title("RAG Document Chatbot (v3)")

# --- MODEL SELECTOR ---
model_option = st.selectbox(
    "Select model",
    [
        "Fast (quick responses, lower accuracy)",
        "Balanced (good speed and quality)",
        "Accurate (slower, best answers)"
    ]
)

def get_model(option):
    if "Fast" in option:
        return "phi"
    elif "Balanced" in option:
        return "mistral"
    else:
        return "mistral-nemo"

LLM_MODEL = get_model(model_option)

# --- INIT MODELS ---
embeddings = OllamaEmbeddings(model=EMBED_MODEL)
reranker = CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2")

# --- SESSION STATE ---
if "db" not in st.session_state:
    st.session_state.db = None

if "uploaded_files" not in st.session_state:
    st.session_state.uploaded_files = set()

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# --- LOAD EXISTING DB ---
def load_existing_vector_store():
    if os.path.exists(VECTOR_STORE_PATH):
        db = FAISS.load_local(
            VECTOR_STORE_PATH,
            embeddings,
            allow_dangerous_deserialization=True
        )
        st.session_state.db = db

load_existing_vector_store()

# --- FILE LOADERS ---
def load_txt(file):
    return file.read().decode("utf-8", errors="ignore")

def load_pdf(file):
    reader = PdfReader(file)
    return "".join([p.extract_text() or "" for p in reader.pages])

def load_docx(file):
    doc = DocxDocument(file)
    return "\n".join([p.text for p in doc.paragraphs])

def load_file(file):
    ext = file.name.split(".")[-1].lower()
    if ext == "txt":
        return load_txt(file)
    elif ext == "pdf":
        return load_pdf(file)
    elif ext == "docx":
        return load_docx(file)
    return None

# --- CHUNKING ---
def build_documents(text, source):
    docs = [Document(page_content=text, metadata={"source": source})]
    splitter = RecursiveCharacterTextSplitter(chunk_size=250, chunk_overlap=30)
    return splitter.split_documents(docs)

# --- PROCESS FILE ---
def process_uploaded_file(file):
    file.seek(0)
    text = load_file(file)

    if not text:
        st.error(f"Failed to extract text from {file.name}")
        return

    chunks = build_documents(text, file.name)

    if os.path.exists(VECTOR_STORE_PATH):
        db = FAISS.load_local(
            VECTOR_STORE_PATH,
            embeddings,
            allow_dangerous_deserialization=True
        )
        db.add_documents(chunks)
    else:
        db = FAISS.from_documents(chunks, embeddings)

    db.save_local(VECTOR_STORE_PATH)
    st.session_state.db = db

# --- RERANK ---
def rerank(query, docs):
    pairs = [(query, d.page_content) for d in docs]
    scores = reranker.predict(pairs)
    ranked = sorted(zip(scores, docs), key=lambda x: x[0], reverse=True)
    return [doc for _, doc in ranked[:2]]

# --- ANSWER ---
def answer_query(query):
    t1 = time.time()

    docs = st.session_state.db.similarity_search(query, k=6)
    docs = rerank(query, docs)

    t2 = time.time()

    context = "\n\n".join(
        f"[Source: {d.metadata.get('source')}]\n{d.page_content}"
        for d in docs
    )

    history = "\n".join(
        f"User: {h['q']}\nAssistant: {h['a']}"
        for h in st.session_state.chat_history[-3:]
    )

    llm = OllamaLLM(model=LLM_MODEL)

    prompt = f"""Use the context and conversation history to answer.

Context:
{context}

Conversation:
{history}

Question: {query}

Answer clearly and concisely:
"""

    response = llm.invoke(prompt).strip()

    t3 = time.time()

    return response, docs, t2 - t1, t3 - t2

# --- UPLOAD UI ---
uploaded_files = st.file_uploader(
    "Upload documents",
    type=None,
    accept_multiple_files=True
)

allowed = ["txt", "pdf", "docx"]

if uploaded_files:
    for file in uploaded_files:
        ext = file.name.split(".")[-1].lower()
        if ext not in allowed:
            st.warning(f"{file.name} not supported")
            continue

        if file.name not in st.session_state.uploaded_files:
            process_uploaded_file(file)
            st.session_state.uploaded_files.add(file.name)
            st.success(f"Processed: {file.name}")

# --- CLEAR DB ---
if st.button("Clear knowledge base"):
    if os.path.exists(VECTOR_STORE_PATH):
        shutil.rmtree(VECTOR_STORE_PATH)

    st.session_state.db = None
    st.session_state.uploaded_files = set()
    st.session_state.chat_history = []

    st.success("Knowledge base cleared")

# --- CHAT ---
with st.form("chat_form"):
    query = st.text_input("Ask a question")
    submit = st.form_submit_button("Submit")

if submit:
    if not st.session_state.db:
        st.warning("Upload documents first")
    elif not query.strip():
        st.warning("Enter a question")
    else:
        response, docs, rt, lt = answer_query(query)

        st.session_state.chat_history.append({
            "q": query,
            "a": response
        })

        st.write(f"Retrieval: {rt:.2f}s | LLM: {lt:.2f}s")

        st.subheader("Answer")
        st.write(response)

        st.markdown("### Sources")
        for i, doc in enumerate(docs[:2]):
            st.markdown(f"**Source {i+1}: {doc.metadata.get('source')}**")
            st.code(doc.page_content[:300])

# --- HISTORY ---
for chat in st.session_state.chat_history:
    st.markdown(f"**User:** {chat['q']}")
    st.markdown(f"**Assistant:** {chat['a']}")